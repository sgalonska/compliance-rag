import os
import pandas as pd
from typing import List, Dict, Any
from pathlib import Path
import PyPDF2
from docx import Document
from openpyxl import load_workbook
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.schema import Document as LangchainDocument

class DocumentProcessor:
    def __init__(self, chunk_size: int = 1000, chunk_overlap: int = 200):
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=chunk_size,
            chunk_overlap=chunk_overlap,
            length_function=len,
        )
    
    def extract_text_from_pdf(self, file_path: str) -> str:
        text = ""
        with open(file_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            for page in pdf_reader.pages:
                text += page.extract_text() + "\n"
        return text
    
    def extract_text_from_docx(self, file_path: str) -> str:
        doc = Document(file_path)
        text = ""
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
        return text
    
    def extract_text_from_excel(self, file_path: str) -> str:
        workbook = load_workbook(file_path, data_only=True)
        text = ""
        
        for sheet_name in workbook.sheetnames:
            sheet = workbook[sheet_name]
            text += f"Sheet: {sheet_name}\n"
            
            for row in sheet.iter_rows(values_only=True):
                row_text = " | ".join([str(cell) if cell is not None else "" for cell in row])
                if row_text.strip():
                    text += row_text + "\n"
            text += "\n"
        
        return text
    
    def process_file(self, file_path: str) -> List[LangchainDocument]:
        file_path = Path(file_path)
        file_extension = file_path.suffix.lower()
        
        if file_extension == '.pdf':
            text = self.extract_text_from_pdf(str(file_path))
        elif file_extension == '.docx':
            text = self.extract_text_from_docx(str(file_path))
        elif file_extension in ['.xlsx', '.xls']:
            text = self.extract_text_from_excel(str(file_path))
        else:
            raise ValueError(f"Unsupported file format: {file_extension}")
        
        chunks = self.text_splitter.split_text(text)
        
        documents = []
        for i, chunk in enumerate(chunks):
            doc = LangchainDocument(
                page_content=chunk,
                metadata={
                    "source": str(file_path),
                    "chunk_id": i,
                    "file_type": file_extension,
                    "file_name": file_path.name
                }
            )
            documents.append(doc)
        
        return documents
    
    def process_directory(self, directory_path: str) -> List[LangchainDocument]:
        supported_extensions = ['.pdf', '.docx', '.xlsx', '.xls']
        all_documents = []
        
        directory = Path(directory_path)
        for file_path in directory.rglob('*'):
            if file_path.suffix.lower() in supported_extensions:
                try:
                    documents = self.process_file(str(file_path))
                    all_documents.extend(documents)
                    print(f"Processed: {file_path.name} ({len(documents)} chunks)")
                except Exception as e:
                    print(f"Error processing {file_path.name}: {str(e)}")
        
        return all_documents